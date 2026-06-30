from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


PROJECT_ROOT = Path("X:/Coding/Statbirt")
REPORT_DIR = PROJECT_ROOT / "reports"
DOCX_PATH = REPORT_DIR / "learned_probability_cohort_report_20260525.docx"
PDF_PATH = REPORT_DIR / "learned_probability_cohort_report_20260525.pdf"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_GRAY = "F2F4F7"
BORDER = "D9E2EC"
CALLOUT = "F4F6F9"
BODY = RGBColor(31, 31, 31)
MUTED = RGBColor(89, 89, 89)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")


def set_keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def style_run(run, *, bold=False, italic=False, color: RGBColor | None = None, size: float | None = None) -> None:
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    set_keep_with_next(paragraph)
    return paragraph


def add_body_paragraph(doc: Document, text: str = ""):
    paragraph = doc.add_paragraph(style="Normal")
    if text:
        paragraph.add_run(text)
    return paragraph


def add_callout(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [9120], indent_dxa=120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    set_cell_border(cell, "D7DEE8")
    set_cell_margins(cell, 140, 180, 140, 180)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    label_run = paragraph.add_run(f"{label}: ")
    style_run(label_run, bold=True, color=RGBColor(31, 58, 95))
    text_run = paragraph.add_run(text)
    style_run(text_run, color=BODY)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_dxa: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table, widths_dxa, indent_dxa=120)

    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        cell = header_cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_border(cell, "C9D4E2")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        style_run(run, bold=True, color=RGBColor(32, 55, 100), size=9.5)

    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cell = cells[idx]
            set_cell_border(cell, "D9E2EC", "4")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(value)
            style_run(run, color=BODY, size=9.2)

    doc.add_paragraph()


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Statbirt learned model analysis")
    style_run(run, color=MUTED, size=8.5)


def build_report() -> Path:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    add_footer(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Learned Model Probability Cohort Analysis")
    style_run(run, bold=True, color=RGBColor(11, 37, 69), size=22)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run("Statbirt | Historical final decisions through 2026-05-22 | Prepared 2026-05-25")
    style_run(r, color=MUTED, size=10.5)

    add_callout(
        doc,
        "Headline",
        "The search found several learned-probability 75-80% cohorts with at least 100 decisions and hit rates above 90%. "
        "The strongest verified slice was 94 hits in 101 decisions, a 93.1% hit rate.",
    )

    add_heading(doc, "Source And Baseline", 1)
    add_body_paragraph(
        doc,
        "The analysis joined learned model probabilities from model_predictions.csv to candidate-side final results in "
        "statbirt_candidates.csv on date, player_id, and game_pk. Only final decisions with a resolved hit/no-hit result were counted.",
    )
    add_table(
        doc,
        ["Population", "Hits / Decisions", "Hit Rate"],
        [
            ["All final candidate decisions", "32,678 / 57,750", "56.6%"],
            ["All learned 75-80% probability picks", "329 / 432", "76.2%"],
        ],
        [5200, 2200, 1960],
    )

    add_heading(doc, "Strict 75-80% Cohorts Above 90%", 1)
    add_table(
        doc,
        ["Cohort", "Hits / Decisions", "Hit Rate"],
        [
            [
                "hitter_ba_25_ab <= .300 + pitcher_hpi_200 >= 1.00 + pitcher_hpi_season >= 1.10",
                "94 / 101",
                "93.1%",
            ],
            [
                "hitter_bb_rate_season >= .080 + hitter_last_5_games_ba <= .400 + pitcher_lr_opp_ba_50 >= .260",
                "103 / 113",
                "91.2%",
            ],
            [
                "hitter_ba_25_ab <= .350 + hitter_bb_rate_season >= .080 + pitcher_lr_opp_ba_50 >= .260",
                "91 / 100",
                "91.0%",
            ],
            [
                "hitter_ba_25_ab <= .300 + pitcher_hits_last_18_ip >= 18 + pitcher_lr_opp_ba_200 >= .260",
                "96 / 106",
                "90.6%",
            ],
            [
                "hitter_ba_2500_ab >= .280 + hitter_bb_rate_500_pa >= .060 + pitcher_hpi_season >= 1.10",
                "101 / 112",
                "90.2%",
            ],
        ],
        [6100, 1760, 1500],
    )

    add_heading(doc, "Original 51/54 Signal", 1)
    add_body_paragraph(
        doc,
        "The original high-signal slice also verified: learned probability of 75-80%, congregation member, and exactly 1-2 stop valves "
        "produced 51 hits in 54 resolved decisions, a 94.4% hit rate. It did not broaden cleanly by itself.",
    )
    add_table(
        doc,
        ["Expansion", "Hits / Decisions", "Hit Rate"],
        [
            ["Original: congregation + exactly 1-2 stop valves", "51 / 54", "94.4%"],
            ["Add 0-stop players", "56 / 60", "93.3%"],
            ["Add 3-stop players", "73 / 82", "89.0%"],
            ["All players, exactly 1-2 stop valves", "85 / 104", "81.7%"],
            ["Congregation, 72-80% probability, exactly 1-2 stop valves", "142 / 168", "84.5%"],
        ],
        [6100, 1760, 1500],
    )

    add_heading(doc, "Interpretation", 1)
    p = add_body_paragraph(doc)
    p.add_run("The strongest broadened signal appears to be: ").bold = True
    p.add_run("good learned probability + recently under-hot hitter + vulnerable starter.")
    add_body_paragraph(
        doc,
        "The best cohort is especially interesting because hitter_ba_25_ab <= .300 sounds counterintuitive at first. "
        "Paired with weak starter HPI, it may be catching good hitters who are not overheated but are in favorable contact matchups.",
    )

    add_heading(doc, "Caveat", 1)
    add_callout(
        doc,
        "Use carefully",
        "These cohorts were mined from historical data, so they should be treated as high-priority hypotheses rather than guaranteed edges. "
        "The 94/101 cohort has a Wilson 95% interval of roughly 86.4% to 96.6%. The next useful step would be walk-forward testing.",
    )

    doc.save(DOCX_PATH)
    return DOCX_PATH


def _pdf_styles():
    styles = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "ReportTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=20,
            leading=24,
            textColor=colors.HexColor("#0B2545"),
            alignment=TA_LEFT,
            spaceAfter=4,
        ),
        "subtitle": ParagraphStyle(
            "ReportSubtitle",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=9.5,
            leading=12,
            textColor=colors.HexColor("#595959"),
            spaceAfter=14,
        ),
        "h1": ParagraphStyle(
            "ReportHeading1",
            parent=styles["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=14,
            leading=17,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=12,
            spaceAfter=6,
        ),
        "body": ParagraphStyle(
            "ReportBody",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.2,
            leading=13.2,
            textColor=colors.HexColor("#1F1F1F"),
            spaceAfter=7,
        ),
        "small": ParagraphStyle(
            "ReportSmall",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=8.6,
            leading=10.8,
            textColor=colors.HexColor("#1F1F1F"),
        ),
        "small_center": ParagraphStyle(
            "ReportSmallCenter",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=8.8,
            leading=10.8,
            textColor=colors.HexColor("#1F1F1F"),
            alignment=TA_CENTER,
        ),
        "table_header": ParagraphStyle(
            "ReportTableHeader",
            parent=styles["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=8.8,
            leading=10.5,
            textColor=colors.HexColor("#203764"),
            alignment=TA_CENTER,
        ),
        "callout": ParagraphStyle(
            "ReportCallout",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10,
            leading=13,
            textColor=colors.HexColor("#1F1F1F"),
            leftIndent=0,
            rightIndent=0,
        ),
    }


def _p(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(text, style)


def _pdf_table(headers: list[str], rows: list[list[str]], col_widths: list[float], styles_map: dict) -> Table:
    data = [[_p(header, styles_map["table_header"]) for header in headers]]
    for row in rows:
        data.append(
            [
                _p(value, styles_map["small"] if idx == 0 else styles_map["small_center"])
                for idx, value in enumerate(row)
            ]
        )
    table = Table(data, colWidths=col_widths, hAlign="LEFT", repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#D9E2EC")),
                ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#C9D4E2")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return table


def _callout(text: str, styles_map: dict) -> Table:
    table = Table(
        [[_p(text, styles_map["callout"])]],
        colWidths=[6.35 * inch],
        hAlign="LEFT",
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F4F6F9")),
                ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor("#D7DEE8")),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 9),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 9),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )
    return table


def build_pdf_report() -> Path:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    styles_map = _pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.72 * inch,
        bottomMargin=0.72 * inch,
        title="Learned Model Probability Cohort Analysis",
        author="Statbirt",
    )
    story = [
        _p("Learned Model Probability Cohort Analysis", styles_map["title"]),
        _p("Statbirt | Historical final decisions through 2026-05-22 | Prepared 2026-05-25", styles_map["subtitle"]),
        _callout(
            "<b>Headline:</b> The search found several learned-probability 75-80% cohorts with at least 100 decisions "
            "and hit rates above 90%. The strongest verified slice was 94 hits in 101 decisions, a 93.1% hit rate.",
            styles_map,
        ),
        Spacer(1, 12),
        _p("Source And Baseline", styles_map["h1"]),
        _p(
            "The analysis joined learned model probabilities from <b>model_predictions.csv</b> to candidate-side final results "
            "in <b>statbirt_candidates.csv</b> on date, player_id, and game_pk. Only final decisions with a resolved hit/no-hit "
            "result were counted.",
            styles_map["body"],
        ),
        _pdf_table(
            ["Population", "Hits / Decisions", "Hit Rate"],
            [
                ["All final candidate decisions", "32,678 / 57,750", "56.6%"],
                ["All learned 75-80% probability picks", "329 / 432", "76.2%"],
            ],
            [3.7 * inch, 1.5 * inch, 1.1 * inch],
            styles_map,
        ),
        Spacer(1, 10),
        _p("Strict 75-80% Cohorts Above 90%", styles_map["h1"]),
        _pdf_table(
            ["Cohort", "Hits / Decisions", "Hit Rate"],
            [
                [
                    "hitter_ba_25_ab <= .300 + pitcher_hpi_200 >= 1.00 + pitcher_hpi_season >= 1.10",
                    "94 / 101",
                    "93.1%",
                ],
                [
                    "hitter_bb_rate_season >= .080 + hitter_last_5_games_ba <= .400 + pitcher_lr_opp_ba_50 >= .260",
                    "103 / 113",
                    "91.2%",
                ],
                [
                    "hitter_ba_25_ab <= .350 + hitter_bb_rate_season >= .080 + pitcher_lr_opp_ba_50 >= .260",
                    "91 / 100",
                    "91.0%",
                ],
                [
                    "hitter_ba_25_ab <= .300 + pitcher_hits_last_18_ip >= 18 + pitcher_lr_opp_ba_200 >= .260",
                    "96 / 106",
                    "90.6%",
                ],
                [
                    "hitter_ba_2500_ab >= .280 + hitter_bb_rate_500_pa >= .060 + pitcher_hpi_season >= 1.10",
                    "101 / 112",
                    "90.2%",
                ],
            ],
            [4.45 * inch, 1.1 * inch, 0.75 * inch],
            styles_map,
        ),
        Spacer(1, 10),
        _p("Original 51/54 Signal", styles_map["h1"]),
        _p(
            "The original high-signal slice also verified: learned probability of 75-80%, congregation member, and exactly "
            "1-2 stop valves produced 51 hits in 54 resolved decisions, a 94.4% hit rate. It did not broaden cleanly by itself.",
            styles_map["body"],
        ),
        _pdf_table(
            ["Expansion", "Hits / Decisions", "Hit Rate"],
            [
                ["Original: congregation + exactly 1-2 stop valves", "51 / 54", "94.4%"],
                ["Add 0-stop players", "56 / 60", "93.3%"],
                ["Add 3-stop players", "73 / 82", "89.0%"],
                ["All players, exactly 1-2 stop valves", "85 / 104", "81.7%"],
                ["Congregation, 72-80% probability, exactly 1-2 stop valves", "142 / 168", "84.5%"],
            ],
            [4.45 * inch, 1.1 * inch, 0.75 * inch],
            styles_map,
        ),
        Spacer(1, 10),
        _p("Interpretation", styles_map["h1"]),
        _p(
            "The strongest broadened signal appears to be: <b>good learned probability + recently under-hot hitter + vulnerable starter.</b>",
            styles_map["body"],
        ),
        _p(
            "The best cohort is especially interesting because <b>hitter_ba_25_ab <= .300</b> sounds counterintuitive at first. "
            "Paired with weak starter HPI, it may be catching good hitters who are not overheated but are in favorable contact matchups.",
            styles_map["body"],
        ),
        _p("Caveat", styles_map["h1"]),
        _callout(
            "<b>Use carefully:</b> These cohorts were mined from historical data, so they should be treated as high-priority "
            "hypotheses rather than guaranteed edges. The 94/101 cohort has a Wilson 95% interval of roughly 86.4% to 96.6%. "
            "The next useful step would be walk-forward testing.",
            styles_map,
        ),
    ]
    doc.build(story)
    return PDF_PATH


if __name__ == "__main__":
    print(build_report())
    print(build_pdf_report())
